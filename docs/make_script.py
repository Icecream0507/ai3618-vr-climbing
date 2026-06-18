#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate the 10-minute, 5-speaker presentation script as a .docx.
Run: python docs/make_script.py  ->  docs/SummitVR_演讲稿.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
RED = RGBColor(0xC8, 0x10, 0x2E)
REDD = RGBColor(0x9E, 0x0B, 0x22)
MUT = RGBColor(0x6B, 0x72, 0x80)
INK = RGBColor(0x1C, 0x1F, 0x26)

doc = Document()

# base font (incl. East Asian)
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_ea(run, font="Microsoft YaHei"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), font)

def para(text="", size=11, color=INK, bold=False, align=None, after=6, before=0, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.4
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
        r.font.bold = bold; r.font.italic = italic; set_ea(r)
    return p

def rich(parts, size=11, after=6, line=1.4):
    """parts: list of (text, bold, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for (t, b, c) in parts:
        r = p.add_run(t); r.font.size = Pt(size); r.font.bold = b; r.font.color.rgb = c
        set_ea(r)
    return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pb.makeelement(qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "12", qn("w:space"): "1", qn("w:color"): "C8102E"})
    pb.append(bottom); pPr.append(pb)

print("helpers ready")

# ========== TITLE BLOCK ==========
para("Summit VR — 项目汇报 演讲稿", 22, RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("课程 AI3618 · 上海交通大学 · 约 10–12 分钟 · 五人分工汇报", 12, MUT, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para("一句话主题：在只有头 + 两手柄的普通 VR 上，把真实抱石的「身体平衡 + 脚法落点」做成可玩机制。",
     11, INK, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
hr()

# ========== HOW TO USE ==========
para("分工与节奏总览", 15, REDD, bold=True, before=6, after=4)
para("全程约 10 分钟，五人接力。每人到自己负责的幻灯片时接讲，讲完一句话过渡给下一位。"
     "Demo 视频在第 8、9 页，播放时讲解者配合旁白。建议各自把自己那段读熟、掐表，留 1 分钟 Q&A 缓冲。",
     11, INK, after=6)

# overview table
rows = [
  ("讲者","负责幻灯片","主题","时长"),
  ("P5 孙艺豪（主讲/开场）","封面 · 目录 · 背景问题","开场 + 我们要解决什么","~1.5 min"),
  ("P1 薛俊智（攀爬系统）","点子 · 调研 · 架构 · 平衡机制","核心技术：头当重心的平衡模型","~2.5 min"),
  ("P2 吴一轩（玩法规则）","脚法常量 · 线路库","脚法抽象 + 5 条线路设计","~2 min"),
  ("P3 邹沛霖（场景美术）","Demo · 化身演进","演示视频 + 画面打磨","~1.5 min"),
  ("P4 陶锐（XR/工程）","整体思路 · 评测体系 · 总结","开发方法论 + 交付总结 + 收尾","~2 min"),
]
table = doc.add_table(rows=len(rows), cols=4)
table.style = "Table Grid"
widths = [Inches(2.0), Inches(2.2), Inches(2.4), Inches(0.9)]
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = table.cell(ri, ci)
        cell.width = widths[ci]
        cell.text = ""
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val); r.font.size = Pt(10); set_ea(r)
        if ri == 0:
            r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"),
                {qn("w:val"):"clear", qn("w:fill"):"C8102E"})
            cell._tc.get_or_add_tcPr().append(shd)
        else:
            r.font.color.rgb = INK
            if ci == 0: r.font.bold = True
para("", after=4)
hr()
print("title + overview done")

def part_header(num, speaker, slides, dur):
    para("", after=2)
    rich([("第 %s 部分　" % num, True, RED), (speaker, True, REDD)], size=16, after=2)
    rich([("对应幻灯片：", False, MUT), (slides, True, INK),
          ("　｜　时长：", False, MUT), (dur, True, INK)], size=10, after=4)

def script_line(text):
    """The actual words to say."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(text); r.font.size = Pt(11.5); r.font.color.rgb = INK; set_ea(r)
    return p

def cue(text):
    """Stage direction / handoff cue."""
    rich([("【提示】", True, RED), (text, False, MUT)], size=10, after=8)

# ===================== PART 1 — P5 =====================
part_header("一", "P5 孙艺豪（主讲 · 开场）", "封面 / 目录 / 背景与问题", "约 1.5 分钟")
cue("站定，等封面页投出再开口。语速放慢，笑一下，这是全场第一印象。")
script_line("各位老师、同学好。先问大家一个问题：你们玩过 VR 攀岩吗？戴上头显，"
            "伸手抓住岩壁上的点，往上一拉，整个人就「爬」上去了——是不是还挺爽的？")
script_line("我们这组的项目也是 VR 攀岩，叫 Summit VR。但在做之前，我们发现现在市面上的 VR 攀岩"
            "游戏——像很火的 The Climb、Gorilla Tag——其实都有同一个毛病：它们全靠两只手。"
            "你就是不停地抓、不停地拉，掉下来也只有两种可能：要么手松了，要么手不够长够不着。")
script_line("可是你去问任何一个真正爬过岩的人，他都会告诉你：攀岩真正难的地方，根本不是手劲大不大，"
            "而是脚怎么踩、重心怎么放。打个比方，你贴在墙上，两只手都抓在右边，身体就会像一扇门一样"
            "「啪」地往外甩开——攀岩里这叫「开门」。怎么破？踩一只左脚顶住，人就稳了。"
            "这种「用脚和重心保持平衡」的功夫，才是攀岩的精髓。")
script_line("而现在的 VR 攀岩，恰恰把这最精髓的一层——脚和平衡——全给砍掉了。"
            "这就是我们想补回来的东西：在一台普普通通、只有头显加两个手柄的 VR 设备上，"
            "把「平衡」和「脚法」做成真正会影响你输赢的玩法。")
script_line("接下来，请负责攀爬核心的薛俊智，给大家讲讲我们到底是怎么把这件事做出来的。")
cue("说完转头示意 P1。目录页可一带而过，或在结尾顺势点一句「整个汇报分六块」。")

# ===================== PART 2 — P1 =====================
part_header("二", "P1 薛俊智（攀爬系统）", "我们的点子 / 调研支撑 / 系统架构 / 核心平衡机制", "约 2.5 分钟")
cue("这是技术核心，最可能被追问。平衡机制那页务必放慢，用「门」「秤」的比喻把它讲透。")
script_line("谢谢艺豪。我先讲我们整个的思路，其实就一句大白话：测不到的，就别硬测，把它「估」出来；"
            "只去做真正影响玩法的那部分。")
script_line("为什么这么说？因为普通 VR 设备只能知道你的头和两只手在哪，它根本看不到你的脚。"
            "那脚怎么办？我们想了三招。第一招，拿头当「重心」。你想，人贴在墙上的时候，"
            "脑袋基本是跟着身体一起晃的，头往哪边歪，重心就往哪边跑，所以盯着头的左右位置，"
            "就能大概知道你有没有失去平衡。")
script_line("第二招，脚不画出来，只算它的「作用」。我们不去做一双逼真的虚拟腿——那种腿在攀岩这种"
            "高抬腿、侧身的姿势下特别容易穿模、特别假。我们只让一只「看不见的脚」自动踩到你身体"
            "下方最近的脚点上，它唯一的作用就是帮你「把站得更稳」。第三招，失衡不是「啪」一下摔死，"
            "而是有一根会慢慢掉的平衡条，给你反应时间，让你能救回来。")
script_line("这套设计也不是我们拍脑袋想的。有两篇论文给我们撑腰：一篇 2020 年的研究专门做实验发现，"
            "在 VR 攀岩里「让你看到脚」比「让你看到手」更能带来真实感；另一篇 2026 年的研究，"
            "就是用「头当重心」来判断人会不会摔，而且做了真人实验验证过。我们等于是把这两个结论，"
            "搬到了人人都买得起的硬件上。")
script_line("讲讲怎么搭的。这里有个很重要的设计叫「加法」：手抓、往上拉这套基本操作是成熟的、稳的，"
            "我们的平衡和脚法是「加」在它上面的两层。你把这两层关掉，游戏立马变回一个普通的纯手攀岩，"
            "照样能玩。所以我们加新东西，永远不会把原来能跑的部分搞坏——这点对做项目很关键。")
script_line("那平衡到底怎么算的？说穿了特别简单，就像一杆秤。我把你所有的「支撑点」——抓住的手、"
            "踩住的脚——在墙的左右方向上拉出一条线段，这就是你的「安全区间」。然后看你的头"
            "（也就是重心）落在哪：落在区间里，稳；歪出区间了，平衡条就开始往下掉，掉光了你就脱手摔下去。")
script_line("现在就能回答那个核心问题了：为什么踩一只脚就能救回来？因为你多踩一个点，"
            "那条「安全区间」就被撑宽了，你原来歪在外面的重心，一下就重新被框进了安全区里。"
            "演示里失衡变红、踩脚回正的那一瞬间，就是我们整个项目最得意的一下。")
script_line("脚和线路具体怎么设计的，请吴一轩接着讲。")
cue("「最得意的一下」说完停半拍，再过渡。这句是全场记忆点。")
print("part 1+2 done")

# ===================== PART 3 — P2 =====================
part_header("三", "P2 吴一轩（玩法规则）", "脚法抽象与关键常量 / 线路库", "约 2 分钟")
cue("线路页有视频缩略图，指着讲。重点把「颜色只是建议」和 The Gap 的梗讲出来，可以轻松点。")
script_line("谢谢俊智。我接着把脚和线路讲细一点。刚才说脚是「看不见」的，那它怎么知道该踩哪？"
            "其实很像真人爬岩：你踩稳一只脚，是不会随便挪的，非得等爬高了、这只脚实在别扭了，"
            "才会换到下一个点。我们就照这个来——脚一旦踩上某个点，就「黏」在那儿，"
            "直到你爬过去够不着了，它才松开去找下一个。这样平衡才稳，不会一直跳来跳去。")
script_line("可能有人会问，干嘛不老老实实做一双会动的腿？因为在攀岩这种高抬腿、侧身的姿势下，"
            "自动算出来的腿十有八九是歪的、穿模的，特别别扭。而我们根本不关心腿长啥样，"
            "只关心「这只脚撑不撑得住」——做双假腿纯属费力不讨好。")
script_line("还有个细节挺有意思：手是有长度的，大概就一臂展，0.88 米左右。够不着的点你就是抓不到，"
            "硬伸也没用，得先把身体挪过去。这一下子就有攀岩那味儿了——不是无脑乱抓，是要规划路线的。")
script_line("线路我们做了 5 条，全是程序自动生成的，不用美术也能跑。其中 4 条能爬到顶，难度一条比一条高："
            "热身的 Warm-up；专门逼你用脚保平衡的 Balance Test；放了「易碎点」、逼你别磨蹭的 The Arête；"
            "还有考验体力、得省着劲爬的 Endurance。")
script_line("这里要澄清一个点：岩点的颜色——黄的、橙的、紫的——现在只是「建议」，不是硬规定。"
            "跟真岩壁一样，你手脚想踩哪个都行。但我们故意把默认线路设计成「纯用手过不去」："
            "你两只手都搬到右边的时候，重心一定会甩出去，这时候你非踩个左脚不可。逼你用脚，懂了吧。")
script_line("最后第 5 条线路最皮，叫 The Gap，是我们故意做成「爬不上去」的——中间空了 2 米，"
            "你怎么伸都够不着。它存在的唯一意义，就是向大家证明：我们这个「手够不够得着」是真的有用的，"
            "不是摆设。好，下面请邹沛霖带大家看实际效果。")
cue("说到「看实际效果」把场子交给 P3，由他控制视频播放。")

# ===================== PART 4 — P3 =====================
part_header("四", "P3 邹沛霖（场景美术 · Demo）", "Demo 核心高光 / 化身演进", "约 1.5 分钟")
cue("先点开第 8 页 demo 视频，边播边讲，旁白卡着画面。视频约 25 秒，别抢拍。")
script_line("好，咱们直接看。大家盯着屏幕上这根平衡条——看，现在这个小人两只手都伸到右边去了，"
            "平衡条开始变红、闪起来了，这就是在告诉你：重心歪出去了，要掉了！")
script_line("注意看这一下——他踩上了左边这个橙色的脚点，平衡条「唰」地就回来了，人稳住了。"
            "就这一下，「变红、踩脚、回正」，是现在市面上所有 VR 攀岩都做不到的。这就是我们的招牌动作。")
script_line("旁边这几个颜色是图例：黄的是手点、橙的是脚点、绿的是终点。要是真掉下去了，"
            "会重生回到检查点，爬到顶的时候，上面会蹦出「Summit」和你的通关时间。")
cue("切第 9 页线路视频，快速扫四条完攀 + The Gap，再切第 11 页化身演进对比图。")
script_line("这几段是四条线路从头爬到顶的完整录像，最后这条够不着的，就是刚才一轩说的 The Gap。")
script_line("再说说画面。我们的小人一开始特别丑，是个直挺挺的假人，跟提线木偶似的。"
            "后来我们找了一款做得很好的物理攀岩游戏当参照，一帧一帧对着改，改成了现在这样——"
            "整个身子是「挂」在抓着的那只手上的，单手吊着时身体会自然倒向那只手，脑袋还一直冲着岩壁看，"
            "一下就像真人在爬了。这些纯为好看，不影响里面的玩法逻辑。")
script_line("接下来请陶锐，给大家讲讲我们整个项目是怎么一步步做出来的，最后交付了些什么。")

# ===================== PART 5 — P4 =====================
part_header("五", "P4 陶锐（XR / 工程 · 收尾）", "整体任务思路 / 评测体系 / 项目总结", "约 2 分钟")
cue("这是收尾，语气要稳、要收得住。结尾「谢谢」后停住，准备接 Q&A。")
script_line("谢谢沛霖。我来说说我们这个项目整体是怎么推进的，以及最后到底做出了什么。")
script_line("我们有个原则，叫「先保证能验证，再谈别的」。我们没一上来就埋头调手感、做美术，"
            "而是先写了个「机器人」——让程序自己当玩家，自动地抓、爬、故意失衡，"
            "然后检查：失衡会不会掉？掉了会不会重生？到顶会不会计时？我们定了 10 条这样的检查，"
            "再加 9 条平衡算法的数学验证，全部一次通过。")
script_line("这么做有个特别实在的好处：核心玩法对不对，机器自动帮我们盯着。所以我们五个人能各干各的、"
            "同时往前推，不用排队抢那一个头显。而且每次有人改了代码，机器人立马重跑，"
            "保证没把原来好好的东西改坏。")
script_line("再说评测。我们把「怎么测」这件事先备齐了：让 5 到 8 个同学分别玩「纯手版」和「平衡加脚法版」，"
            "记录他们的通关时间、摔了几次、分别是因为啥摔的，再用两份业界标准问卷量一下"
            "「累不累」和「晕不晕」。这套采集的工具和流程我们都做好了。"
            "我们赌的是：加了平衡和脚法之后，玩家会觉得更真实、更有挑战，但不会更晕——"
            "因为难度是动脑子的难度，不是那种让人反胃的晃动。")
script_line("最后总结一下我们交付了啥：平衡机制、脚法、5 条线路、那套自动检查、"
            "屏幕上的血条和音效、能上手玩的两个视角、演示视频、还有报告和答辩材料——全都做完了。"
            "唯一还没做的，就是真的把那 5 到 8 个同学拉来做实验。")
script_line("一句话收尾：Summit VR 证明了一件事——脚法和平衡，真的可以用很低的成本、"
            "很稳妥地加进只有手柄的 VR 攀岩里。谢谢大家，欢迎提问。")
cue("「谢谢大家」后五人一起致意。Q&A 细节见 docs/DEFENSE_QA.md，对应模块的问题由对应负责人接。")
print("parts 3+4+5 done")

# save
out = os.path.join(HERE, "SummitVR_演讲稿.docx")
doc.save(out)
print("SAVED ->", out)
